import docx
import pandas as pd

import re
import pdftotext

def get_clauses_from_word(doc):
    """
    Given a word document, returns the sentences in it for the downstream processing to happen
    Returns a dataframe with 'Clause' as the column
    """
    print("in get_clauses_from_word fun")
    try:
        wdoc = docx.Document(doc)
    except:
        print("Given file doesn't appear to be a word document")
        return None
    print("Doc {} has {} clauses".format(doc, len(wdoc.paragraphs)))

    paras = [p.text for p in wdoc.paragraphs]

    wdocdf = pd.DataFrame(paras, columns=['Clause'])
    nan_value = float("NaN")
    wdocdf.replace("", nan_value, inplace=True)
    wdocdf.dropna(subset=['Clause'], inplace=True)
    wdocdf.reset_index(drop=True, inplace=True)
    
    print("Doc {} has {} clauses after cleaning".format(doc, len(wdocdf)))

    return wdocdf


def get_clauses_from_pdf(pdf):
    print("in get_clauses_from_pdf fun")
    # compile a substitution for page
    page_sub = re.compile('Page\s+\d+\s+of\s+\d+', re.IGNORECASE)

    #with open(pdf, "rb") as f:
    pdf_file = pdftotext.PDF(pdf)

    sentences = []
    for page in pdf_file:
        lines = page.split('\n')
        lines = [nl.replace('\r','') for nl in lines]
        ls_lines = [nl.lstrip() for nl in lines]
        ls_lines = [page_sub.sub(" ", line) for line in ls_lines]
        sent = None
        for n in ls_lines:
            sent = n if sent is None else sent+" "+n
            #sent = sent+" " + n
            if n.endswith('.') or n.endswith(':'):
                sentences.append(sent)
                sent = None
    # sentences now has all the clauses
    df = pd.DataFrame(sentences, columns=['Clause'])

    return df

if __name__ == "__main__":
    clauses = get_clauses_from_pdf('../reference/sample_eula_1.pdf')